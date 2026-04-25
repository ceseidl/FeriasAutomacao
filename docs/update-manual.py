"""
Atualiza docs/MANUAL.docx in-place para refletir as mudancas do v1.0.0:
  1. Reescreve o paragrafo do controle "Ano" pra descrever o year picker
     estilo calendario (substituiu o NumericUpDown).
  2. Adiciona uma regra explicita sobre o filtro estrito por ano.
  3. Substitui a screenshot embedada por docs/screenshot.png atualizado.
  4. Corrige o rodape que ainda fala em MANUAL.md como fonte (a fonte
     virou o proprio .docx).
  5. Adiciona instrucoes de instalacao via MSI (com descricao do dialogo
     de selecao de atalhos) e renumera as subsecoes 1.x.
  6. Atualiza a descricao da aba Ferias pra refletir as 5 abas atuais
     (Ferias + Squads + Integrantes + Status + Instrucoes) e renomeia
     ferias-2026.xlsx -> Ferias-template.xlsx em todo o documento.
  7. Atualiza a saida pra pasta unica configuravel com nomes fixos
     (Ferias.md / Ferias.html / Ferias.xlsx / Ferias.pdf) que sobrescreve
     a execucao anterior. Sem subpastas com timestamp.
  8. Insere descricao do controle "Pasta de saida" da GUI (com
     persistencia em HKCU) logo apos o item do controle "Ano".

Idempotente: rodar mais de uma vez nao duplica conteudo.
"""
import sys
from pathlib import Path
from docx import Document

ROOT = Path(__file__).resolve().parent.parent
DOCX = ROOT / 'docs' / 'MANUAL.docx'
NEW_PNG = ROOT / 'docs' / 'screenshot.png'

doc = Document(str(DOCX))

# ---------- 1) Paragrafo do controle "Ano" ----------
target_old = 'Ano - ano que aparece no titulo do relatorio e no Gantt'
new_text = (
    'Ano - ano que aparece no titulo do relatorio, no Gantt e que filtra '
    'a planilha. Padrao = ano atual. Clicar no botao "ANO ▾" abre um '
    'calendario 4x3 com 12 anos por pagina; usar as setas no topo do '
    'popup para navegar de decada em decada e clicar no ano desejado. '
    'O popup fecha sozinho ao escolher um ano, ao apertar Esc ou ao '
    'clicar fora dele.'
)
hits = 0
for p in doc.paragraphs:
    if p.text.startswith(target_old):
        # zera os runs e reescreve preservando o estilo
        for r in list(p.runs):
            r.text = ''
        if p.runs:
            p.runs[0].text = new_text
        else:
            p.add_run(new_text)
        hits += 1
        break
if hits == 0:
    print('AVISO: paragrafo do controle Ano nao encontrado', file=sys.stderr)

# ---------- 2) Regra do filtro estrito por ano ----------
# Insere apos o paragrafo "2.3 Regras" como mais um item Compact.
regra_filtro = (
    'O relatorio so traz ferias do ano selecionado no picker. Mesmo que '
    'a planilha tenha linhas de outros anos, elas sao filtradas fora '
    'antes de gerar o relatorio.'
)
inserted = False
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'Salvar e fechar a planilha antes de gerar o relatorio (Excel trava o arquivo)':
        # Insere um novo paragrafo logo apos, com o mesmo estilo "Compact".
        new_p = p.insert_paragraph_before('')  # placeholder so pra ter o objeto
        # insert_paragraph_before insere ANTES; entao trocamos a abordagem:
        # cria via XML clonando o atual.
        from copy import deepcopy
        clone = deepcopy(p._element)
        # zera o texto do clone
        for t in clone.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
            t.text = regra_filtro
            break
        # remove os outros <w:t> alem do primeiro pra nao concatenar lixo
        ts = list(clone.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'))
        for extra in ts[1:]:
            extra.getparent().remove(extra)
        p._element.addnext(clone)
        # remove o placeholder vazio que criei sem querer
        new_p._element.getparent().remove(new_p._element)
        inserted = True
        break
if not inserted:
    print('AVISO: regra do filtro nao foi inserida', file=sys.stderr)

# ---------- 3) Rodape ----------
for p in doc.paragraphs:
    if p.text.startswith('Este manual e gerado a cada release a partir de docs/MANUAL.md'):
        for r in list(p.runs):
            r.text = ''
        if p.runs:
            p.runs[0].text = (
                'Este manual e mantido em docs/MANUAL.docx (fonte direta). '
                'Para atualizar, edite o .docx e rode docs/update-manual.py se '
                'precisar reaplicar mudancas em massa.'
            )
        break

# ---------- 4) Substitui a screenshot embedada ----------
# python-docx nao tem API publica pra trocar a imagem; manipulo a parte do pacote.
new_bytes = NEW_PNG.read_bytes()
swapped = False
for rel_id, rel in doc.part.rels.items():
    target = rel.target_ref
    if 'media' in target and target.endswith('.png'):
        # rel.target_part eh o ImagePart
        rel.target_part._blob = new_bytes
        swapped = True
        print(f'Screenshot substituida em {target} ({len(new_bytes)} bytes)')
if not swapped:
    print('AVISO: nenhuma imagem encontrada pra substituir', file=sys.stderr)

# ---------- 5) Instalacao via MSI ----------
# Plano:
#   - Reescreve o paragrafo "Clicar em FeriasAutomacao-latest.zip..." pra
#     listar as DUAS opcoes (MSI recomendado + ZIP portatil).
#   - Insere uma nova secao "1.2 Instalar via MSI (recomendado)" antes
#     da atual "1.2 Extrair o ZIP", com paragrafos Compact descrevendo
#     o fluxo do instalador (incluindo a tela de selecao de atalhos).
#   - Renumera "1.2 Extrair o ZIP" -> "1.3 Alternativa: extrair o ZIP"
#     e "1.3 Conferir o conteudo" -> "1.4 Conferir o conteudo".
#
# Idempotente: se o heading "1.2 Instalar via MSI" ja existir, pula
# tudo; se os renumerados ja existirem, nao reescreve.

from copy import deepcopy
W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

def set_text(p, txt):
    """Reescreve um paragrafo preservando o estilo do primeiro run."""
    for r in list(p.runs):
        r.text = ''
    if p.runs:
        p.runs[0].text = txt
    else:
        p.add_run(txt)

def clone_with_text(template_p, txt):
    """Clona um paragrafo (preservando o estilo dele) e poe o texto novo."""
    clone = deepcopy(template_p._element)
    ts = list(clone.iter(f'{W_NS}t'))
    if ts:
        ts[0].text = txt
        for extra in ts[1:]:
            extra.getparent().remove(extra)
    else:
        # se nao tem <w:t>, cria um run+text minimo dentro do paragrafo
        from docx.oxml.ns import qn
        r = clone.makeelement(qn('w:r'), {})
        t = clone.makeelement(qn('w:t'), {})
        t.text = txt
        r.append(t)
        clone.append(r)
    return clone

# ja foi feito antes? se sim, sai cedo
already_done = any(p.text.strip() == '1.2 Instalar via MSI (recomendado)'
                   for p in doc.paragraphs)

if not already_done:
    # 5.1) reescreve o "Clicar em FeriasAutomacao-latest.zip..."
    for p in doc.paragraphs:
        if p.text.strip().startswith('Clicar em FeriasAutomacao-latest.zip'):
            set_text(p, (
                'Baixar uma das opcoes disponiveis na pagina (recomendado: '
                'FeriasAutomacao-1.0.0.msi para instalacao automatica; '
                'FeriasAutomacao-latest.zip para versao portatil sem '
                'instalar nada).'
            ))
            break

    # 5.2) localiza o "1.2 Extrair o ZIP" pra usar como ponto de insercao
    p_extrair = None
    p_conferir = None
    for p in doc.paragraphs:
        if p.text.strip() == '1.2 Extrair o ZIP':
            p_extrair = p
        elif p.text.strip() == '1.3 Conferir o conteudo':
            p_conferir = p

    if p_extrair is None:
        print('AVISO: heading "1.2 Extrair o ZIP" nao encontrado, nao da pra inserir MSI', file=sys.stderr)
    else:
        # pega um Compact ja existente como template de estilo
        compact_template = None
        for p in doc.paragraphs:
            if p.style and p.style.name == 'Compact' and p.text.strip():
                compact_template = p
                break
        if compact_template is None:
            print('AVISO: nao achei nenhum paragrafo Compact pra usar como template', file=sys.stderr)

        # 5.3) renumera "1.2 Extrair o ZIP" -> "1.3 Alternativa: ..."
        set_text(p_extrair, '1.3 Alternativa: extrair o ZIP')

        # 5.4) renumera "1.3 Conferir o conteudo" -> "1.4 ..."
        if p_conferir is not None:
            set_text(p_conferir, '1.4 Conferir o conteudo')

        # 5.5) insere a nova secao ANTES de p_extrair.
        # ordem de insercao: heading + N paragrafos Compact + paragrafo de
        # transicao. Tudo via .addprevious() pra manter a ordem natural.
        msi_blocks = [
            ('Heading 3', '1.2 Instalar via MSI (recomendado)'),
            ('First Paragraph',
             'O .msi instala o app por usuario (nao precisa de admin) em '
             '%LocalAppData%\\Programs\\FeriasAutomacao e cria os atalhos '
             'no menu Iniciar e na area de trabalho automaticamente.'),
            ('Compact', 'Dar duplo clique em FeriasAutomacao-1.0.0.msi.'),
            ('Compact',
             'O instalador mostra uma tela de selecao de features tipo '
             'arvore. Por padrao tudo vem marcado:'),
            ('Compact',
             '- Aplicativo Ferias Automacao (obrigatorio): copia a planilha, '
             'os scripts, o icone e o instalador offline do Pandoc.'),
            ('Compact',
             '- Atalhos no Menu Iniciar (recomendado): cria a pasta '
             '"Programas > Ferias Automacao" com os atalhos Gerar Relatorio, '
             'Manual e Pasta de instalacao.'),
            ('Compact',
             '- Atalho na Area de Trabalho (recomendado): cria o atalho '
             '"Gerar Relatorio" direto na area de trabalho.'),
            ('Compact',
             '- Atalho em "Enviar para" (opcional, vem desmarcado): '
             'adiciona o app ao menu de clique-direito do Windows. '
             'Marcar so se quiser abrir uma planilha .xlsx direto pelo '
             'Enviar para.'),
            ('Compact',
             'Clicar em Avancar > Instalar e aguardar. Ao terminar, abrir '
             'o atalho Gerar Relatorio (menu Iniciar ou desktop) e seguir '
             'para a secao 2.'),
            ('First Paragraph',
             'Quem usar o MSI pode pular as secoes 1.3 e 1.4 abaixo - elas '
             'sao so para quem prefere a versao portatil em ZIP.'),
        ]

        # acha um paragrafo template pra cada estilo usado
        templates = {}
        for p in doc.paragraphs:
            if p.style and p.text.strip():
                name = p.style.name
                if name in {'Heading 3', 'First Paragraph', 'Compact'} and name not in templates:
                    templates[name] = p
        # checa que temos os 3
        missing = {'Heading 3', 'First Paragraph', 'Compact'} - templates.keys()
        if missing:
            print(f'AVISO: estilos nao encontrados: {missing}', file=sys.stderr)

        for style, txt in msi_blocks:
            tpl = templates.get(style)
            if tpl is None:
                continue
            new_p = clone_with_text(tpl, txt)
            p_extrair._element.addprevious(new_p)

        print('OK: secao "1.2 Instalar via MSI (recomendado)" inserida.')
else:
    print('Secao MSI ja presente; nada a fazer em (5).')

# ---------- 6.1) Atualiza descricao da aba Ferias na secao 2.1 ----------
# Antes mencionava apenas "uma aba chamada Ferias". Agora a planilha tem
# 5 abas (Ferias + Squads + Integrantes + Status + Instrucoes).
old_aba = ('A planilha deve ter uma aba chamada Ferias '
           '(e dela que o app le os dados).')
new_aba = ('A planilha tem 5 abas: Ferias (a principal, onde o app le os dados), '
           'Squads (lista de squads), Integrantes (lista de pessoas com seu squad), '
           'Status (Aprovada/Planejada/Solicitada) e Instrucoes (texto de ajuda). '
           'As 3 abas de listas alimentam os dropdowns das colunas Colaborador, '
           'Squad e Status da aba Ferias - pra adicionar uma pessoa nova ou um '
           'squad novo, edite a aba correspondente.')
for p in doc.paragraphs:
    if p.text.strip() == old_aba:
        for r in list(p.runs):
            r.text = ''
        if p.runs:
            p.runs[0].text = new_aba
        else:
            p.add_run(new_aba)
        break

# ---------- 6.2) Renomeacao da planilha (ferias-2026.xlsx -> Ferias-template.xlsx) ----------
# Substitui o nome em todos os runs de todos os paragrafos. Idempotente: se
# o nome ja foi trocado, nao acha mais e nao faz nada.
old_name = 'ferias-2026.xlsx'
new_name = 'Ferias-template.xlsx'
swaps = 0
for p in doc.paragraphs:
    for r in p.runs:
        if old_name in r.text:
            r.text = r.text.replace(old_name, new_name)
            swaps += 1
print(f'Renomeacao da planilha: {swaps} run(s) atualizados.')

# ---------- 7) Saida em pasta unica com nomes fixos (sobrescreve a anterior) ----------
# REGRA NOVA: sem subpasta com timestamp e sem timestamp no nome. Os arquivos
# vao sempre pra mesma pasta (results\ por padrao, ou outra escolhida pelo
# usuario na GUI), com nomes fixos: Ferias.md, Ferias.html, Ferias.xlsx
# (e Ferias.pdf se a opcao estiver marcada). Cada execucao sobrescreve a
# anterior.
#
# A escolha do usuario na GUI fica persistida em
# HKCU:\Software\FeriasAutomacao\OutputDir, entao na proxima abertura a
# janela ja vem apontando pra pasta preferida.
#
# Atualiza (idempotente, aceitando como entrada o estado atual OU o estado
# antigo "arquivos soltos com timestamp"):
#   - paragrafo descritivo
#   - bloco Source Code com a estrutura de pastas
#   - referencia em "Subir o arquivo ... na biblioteca do SharePoint"

novo_desc = ('Os arquivos sao gerados na pasta de saida configurada na GUI '
             '(padrao: results\\) e cada execucao sobrescreve a anterior. '
             'A escolha do usuario fica persistida pra proxima execucao:')
descs_antigos = {
    'Cada execucao cria arquivos com timestamp na pasta results\\:',
    ('Cada execucao cria uma subpasta results\\Ferias-{timestamp}\\ '
     'com o relatorio em multiplos formatos + a planilha-fonte:'),
}
for p in doc.paragraphs:
    if p.text.strip() in descs_antigos or p.text.strip() == novo_desc:
        for r in list(p.runs):
            r.text = ''
        if p.runs:
            p.runs[0].text = novo_desc
        else:
            p.add_run(novo_desc)
        break

# Bloco Source Code: troca pelo novo layout flat (4 arquivos com nome fixo).
nova_arvore = (
    '<pasta-de-saida>\\\n'
    '  Ferias.md     <- versao Markdown (texto puro)\n'
    '  Ferias.html   <- versao HTML formatada (interativa)\n'
    '  Ferias.xlsx   <- copia da planilha que gerou o relatorio\n'
    '  Ferias.pdf    <- (opcional, com a opcao "Gerar PDF tambem" marcada)'
)
for p in doc.paragraphs:
    if p.style.name == 'Source Code' \
            and (p.text.lstrip().startswith('results\\') or p.text.lstrip().startswith('<pasta-de-saida>\\')) \
            and 'Ferias.html' not in p.text:
        for r in list(p.runs):
            r.text = ''
        if p.runs:
            p.runs[0].text = nova_arvore
        else:
            p.add_run(nova_arvore)
        break

# "Subir o arquivo Ferias.pdf na biblioteca do SharePoint" (sem timestamp).
novo_share = 'Subir o arquivo Ferias.pdf da pasta de saida na biblioteca do SharePoint'
shares_antigos = {
    'Subir o arquivo Ferias-{timestamp}.pdf na biblioteca do SharePoint',
    ('Subir o arquivo Ferias-{timestamp}.pdf de dentro da subpasta '
     'Ferias-{timestamp}\\ na biblioteca do SharePoint'),
}
for p in doc.paragraphs:
    if p.text.strip() in shares_antigos or p.text.strip() == novo_share:
        for r in list(p.runs):
            r.text = ''
        if p.runs:
            p.runs[0].text = novo_share
        else:
            p.add_run(novo_share)
        break

# ---------- 8) Controle "Pasta de saida" + persistencia em HKCU ----------
# Insere um paragrafo Compact descrevendo o novo controle da GUI logo
# apos o item "Ano - ano que aparece no titulo...". Idempotente: pula se
# ja existe.
desc_pasta = (
    'Pasta de saida - pasta onde os arquivos serao gerados (padrao: '
    'results\\). Clicar em Procurar... abre um seletor de pasta; a '
    'escolha fica salva em HKCU:\\Software\\FeriasAutomacao\\OutputDir '
    'pra abrir ja preenchida na proxima execucao. Cada execucao '
    'sobrescreve os arquivos anteriores nessa pasta.'
)
ja_tem = any(p.text.strip().startswith('Pasta de saida - pasta onde')
             for p in doc.paragraphs)
if not ja_tem:
    # acha o paragrafo do "Ano" pra inserir o novo logo depois
    p_ano = None
    for p in doc.paragraphs:
        if p.text.startswith('Ano - ano que aparece no titulo'):
            p_ano = p
            break
    if p_ano is not None:
        # acha um Compact pra usar como template de estilo
        compact_tpl = None
        for p in doc.paragraphs:
            if p.style and p.style.name == 'Compact' and p.text.strip():
                compact_tpl = p
                break
        if compact_tpl is not None:
            new_p = clone_with_text(compact_tpl, desc_pasta)
            p_ano._element.addnext(new_p)
            print('OK: paragrafo "Pasta de saida" inserido apos "Ano".')
        else:
            print('AVISO: nenhum Compact disponivel pra clonar (pasta de saida)', file=sys.stderr)
    else:
        print('AVISO: paragrafo "Ano - ..." nao encontrado, pulando descricao da pasta de saida', file=sys.stderr)
else:
    print('Paragrafo "Pasta de saida" ja presente; nada a fazer em (8).')

doc.save(str(DOCX))
print(f'OK: {DOCX} atualizado.')
