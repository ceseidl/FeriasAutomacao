"""
Hotfix do MANUAL.docx pra v1.0.0 (final, com hotfixes pos-release inicial).

Aplica:
  1. Limpa paragrafos duplicados de "O relatorio so traz ferias..."
     (bug de idempotencia em rodadas anteriores do update-manual.py).
  2. Atualiza secao 1.2 (Instalar via MSI):
     - Substitui descricao do FeatureTree pelo InstallDir (tela dedicada
       pra escolher pasta de instalacao).
     - Renomeia "Programas > Ferias Automacao" pra "AI R" (nova pasta no
       Menu Iniciar).
     - Adiciona descricao do atalho "Desinstalar Ferias Automacao".
     - Adiciona descricao do checkbox "Iniciar Ferias Automacao agora"
       no fim do instalador.
  3. Substitui "icone da palmeira" -> "logo da AI R" (apos troca do
     emoji pelo logo vetorial no relatorio).
  4. Adiciona aviso sobre single-instance lock (apenas uma janela por PC).
  5. Adiciona aviso sobre lock estrito ao template oficial (so aceita
     Ferias-template.xlsx com as 5 abas obrigatorias).
  6. Adiciona troubleshooting de reparo via msiexec /f.

Idempotente: rodar mais de uma vez nao duplica conteudo.

Uso:
    python docs/update-manual-hotfix.py
"""
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
DOCX = ROOT / 'docs' / 'MANUAL.docx'

W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

doc = Document(str(DOCX))


def set_text(p, txt):
    """Reescreve um paragrafo preservando o estilo do primeiro run."""
    for r in list(p.runs):
        r.text = ''
    if p.runs:
        p.runs[0].text = txt
    else:
        p.add_run(txt)


def clone_with_text(template_p, txt):
    """Clona um paragrafo preservando estilo e poe o texto novo."""
    clone = deepcopy(template_p._element)
    ts = list(clone.iter(f'{W_NS}t'))
    if ts:
        ts[0].text = txt
        for extra in ts[1:]:
            extra.getparent().remove(extra)
    else:
        r = clone.makeelement(qn('w:r'), {})
        t = clone.makeelement(qn('w:t'), {})
        t.text = txt
        r.append(t)
        clone.append(r)
    return clone


def find_template(style_name):
    """Devolve o primeiro paragrafo nao-vazio com o estilo dado, pra clonar."""
    for p in doc.paragraphs:
        if p.style and p.style.name == style_name and p.text.strip():
            return p
    return None


# ============================================================
# 1) Limpa duplicados de "O relatorio so traz ferias..."
# ============================================================
target_dup = 'O relatorio so traz ferias do ano selecionado no picker'
encontrados = [p for p in doc.paragraphs if p.text.strip().startswith(target_dup)]
removidos = 0
for p in encontrados[1:]:  # mantem o primeiro, remove os outros
    p._element.getparent().remove(p._element)
    removidos += 1
print(f'(1) Removidos {removidos} paragrafos duplicados (de {len(encontrados)} encontrados).')


# ============================================================
# 2) Atualiza secao 1.2 (Instalar via MSI)
# ============================================================
substitutions_1_2 = {
    'O .msi instala o app por usuario (nao precisa de admin) em %LocalAppData%\\Programs\\FeriasAutomacao e cria os atalhos no menu Iniciar e na area de trabalho automaticamente.':
        'O .msi instala o app por usuario (sem prompt de UAC). A pasta padrao '
        'e %LocalAppData%\\Programs\\FeriasAutomacao, mas o instalador permite '
        'escolher outra pasta. Os atalhos do Menu Iniciar ficam na pasta "AI R" '
        '(submenu Iniciar > AI R) e na area de trabalho.',

    'O instalador mostra uma tela de selecao de features tipo arvore. Por padrao tudo vem marcado:':
        'Fluxo do instalador (5 telas): Welcome -> License Agreement -> Choose '
        'installation folder (botao "Change..." pra trocar a pasta) -> '
        'Verify Ready -> Finish.',

    '- Aplicativo Ferias Automacao (obrigatorio): copia a planilha, os scripts, o icone e o instalador offline do Pandoc.':
        'A pasta "AI R" no Menu Iniciar tem 4 atalhos: Gerar Relatorio, Manual '
        'do usuario, Pasta de instalacao e Desinstalar Ferias Automacao.',

    '- Atalhos no Menu Iniciar (recomendado): cria a pasta "Programas > Ferias Automacao" com os atalhos Gerar Relatorio, Manual e Pasta de instalacao.':
        'Atalho "Desinstalar Ferias Automacao": chama o assistente de remocao '
        'do Windows. Equivalente a abrir Configuracoes > Aplicativos e mandar '
        'desinstalar.',

    '- Atalho na Area de Trabalho (recomendado): cria o atalho "Gerar Relatorio" direto na area de trabalho.':
        'Na tela final do instalador, o checkbox "Iniciar Ferias Automacao '
        'agora" ja vem marcado. Se mantiver marcado e clicar Finish, o app '
        'abre direto. Desmarcar so se preferir abrir manualmente depois pelo '
        'atalho.',

    '- Atalho em "Enviar para" (opcional, vem desmarcado): adiciona o app ao menu de clique-direito do Windows. Marcar so se quiser abrir uma planilha .xlsx direto pelo Enviar para.':
        'Para reabrir depois: Menu Iniciar > AI R > Gerar Relatorio (ou o '
        'atalho na area de trabalho).',

    'Clicar em Avancar > Instalar e aguardar. Ao terminar, abrir o atalho Gerar Relatorio (menu Iniciar ou desktop) e seguir para a secao 2.':
        'Apos a instalacao, seguir para a secao 2 (preencher a planilha) e '
        '3 (gerar o relatorio).',
}

aplicados = 0
for p in doc.paragraphs:
    txt = p.text.strip()
    if txt in substitutions_1_2:
        set_text(p, substitutions_1_2[txt])
        aplicados += 1
print(f'(2) Atualizados {aplicados} paragrafos da secao 1.2 MSI.')


# ============================================================
# 3) Substitui "palmeira" por "logo da AI R"
# ============================================================
trocas_logo = 0
for p in doc.paragraphs:
    for r in p.runs:
        if 'palmeira' in r.text.lower():
            txt = r.text
            txt = txt.replace('icone da palmeira', 'logo da AI R')
            txt = txt.replace('icone de palmeira', 'logo da AI R')
            txt = txt.replace('palmeira', 'logo da AI R')
            r.text = txt
            trocas_logo += 1
print(f'(3) Substituidas {trocas_logo} mencoes a palmeira pelo logo AI R.')


# ============================================================
# 4) Single-instance lock
# ============================================================
ja_tem_single = any('Apenas uma janela do app por PC' in p.text for p in doc.paragraphs)
if not ja_tem_single:
    for p in doc.paragraphs:
        if p.text.strip().startswith('Duplo-clique em Gerar Relatorio'):
            compact_tpl = find_template('Compact')
            if compact_tpl is not None:
                txt = ('Apenas uma janela do app por PC: se ja tiver uma janela aberta e voce '
                       'der duplo-clique no atalho de novo, a janela existente volta pra frente '
                       'em vez de abrir outra. Mensagem mostrada: "Ferias Automacao ja esta em '
                       'execucao."')
                new_p = clone_with_text(compact_tpl, txt)
                p._element.addnext(new_p)
                print('(4) Paragrafo single-instance inserido.')
            break
else:
    print('(4) Single-instance ja documentado, pulando.')


# ============================================================
# 5) Lock estrito ao template oficial
# ============================================================
ja_tem_lock = any('so aceita a planilha-template oficial' in p.text for p in doc.paragraphs)
if not ja_tem_lock:
    for p in doc.paragraphs:
        if p.text.strip() == 'Abrir Ferias-template.xlsx no Excel.':
            block_tpl = find_template('Block Text')
            if block_tpl is not None:
                txt = ('Importante: o app so aceita a planilha-template oficial, com nome '
                       'exato Ferias-template.xlsx e as 5 abas obrigatorias (Ferias, Squads, '
                       'Integrantes, Status, Instrucoes). Renomear o arquivo ou mexer na '
                       'estrutura faz o app rejeitar com mensagem clara antes de gerar o '
                       'relatorio.')
                new_p = clone_with_text(block_tpl, txt)
                p._element.addnext(new_p)
                print('(5) Aviso de lock do template oficial inserido.')
            break
else:
    print('(5) Lock do template ja documentado, pulando.')


# ============================================================
# 6) Reparo MSI no troubleshooting (secao 7)
# ============================================================
ja_tem_reparo = any('msiexec /f' in p.text for p in doc.paragraphs)
if not ja_tem_reparo:
    p_problemas = None
    for p in doc.paragraphs:
        if p.text.strip() == '7. Problemas comuns':
            p_problemas = p
            break
    if p_problemas is not None:
        compact_tpl = find_template('Compact')
        if compact_tpl is not None:
            txt = ('App parou de abrir / arquivo da pasta de instalacao foi deletado: '
                   'rodar de novo o .msi do release (oferece "Repair") ou abrir um console '
                   'e executar msiexec /f "FeriasAutomacao-1.0.0.1.msi" pra reinstalar os '
                   'arquivos faltantes. Como ultima alternativa, desinstalar pelo Painel '
                   'de Controle e reinstalar.')
            new_p = clone_with_text(compact_tpl, txt)
            p_problemas._element.addnext(new_p)
            print('(6) Troubleshooting de reparo MSI inserido.')
else:
    print('(6) Reparo MSI ja documentado, pulando.')


# ============================================================
# 7) Remove mencao ao checkbox "Iniciar agora" (que foi removido do MSI)
# ============================================================
# O update anterior tinha mencionado um checkbox "Iniciar Ferias Automacao
# agora" no fim do instalador. Esse checkbox foi removido depois (4 tentativas
# de fazer ele disparar o app falharam). Agora o usuario abre manualmente
# pelo atalho do Desktop ou Menu Iniciar.
substituicoes_7 = {
    'Na tela final do instalador, o checkbox "Iniciar Ferias Automacao agora" ja vem marcado. Se mantiver marcado e clicar Finish, o app abre direto. Desmarcar so se preferir abrir manualmente depois pelo atalho.':
        'Na tela final, clicar Finish pra concluir a instalacao. Apos o '
        'instalador fechar, abrir o app pelo atalho "Gerar Relatorio" '
        '(Area de Trabalho -- sempre criado -- ou Menu Iniciar > AI R).',
}
trocas_7 = 0
for p in doc.paragraphs:
    txt = p.text.strip()
    if txt in substituicoes_7:
        set_text(p, substituicoes_7[txt])
        trocas_7 += 1
print(f'(7) Substituidas {trocas_7} mencoes ao checkbox removido.')


# ============================================================
# 8) Atualiza o nome do MSI nas referencias (release v1.0.0.1)
# ============================================================
# Substitui FeriasAutomacao-1.0.0.msi -> FeriasAutomacao-1.0.0.1.msi
# em todos os runs de todos os paragrafos. Idempotente: se ja foi
# trocado, nao acha mais.
old_msi = 'FeriasAutomacao-1.0.0.msi'
new_msi = 'FeriasAutomacao-1.0.0.1.msi'
trocas_8 = 0
for p in doc.paragraphs:
    for r in p.runs:
        if old_msi in r.text:
            r.text = r.text.replace(old_msi, new_msi)
            trocas_8 += 1
print(f'(8) Substituidos {trocas_8} run(s) com nome do MSI antigo.')


# ============================================================
# 9) Atualiza descricao da aba Integrantes em 2.1 -- 4 colunas
# ============================================================
old_2_1 = ('A planilha tem 5 abas: Ferias (a principal, onde o app le os dados), '
           'Squads (lista de squads), Integrantes (lista de pessoas com seu squad), '
           'Status (Aprovada/Planejada/Solicitada) e Instrucoes (texto de ajuda). '
           'As 3 abas de listas alimentam os dropdowns das colunas Colaborador, '
           'Squad e Status da aba Ferias - pra adicionar uma pessoa nova ou um '
           'squad novo, edite a aba correspondente.')
new_2_1 = ('A planilha tem 5 abas: Ferias (a principal, onde o app le os dados), '
           'Squads (lista de squads), Integrantes (cadastro de pessoas com 4 colunas: '
           'nome, squad, data de admissao na AIR e data fim das ultimas ferias), '
           'Status (Aprovada/Planejada/Solicitada) e Instrucoes (texto de ajuda). '
           'As abas de listas alimentam os dropdowns das colunas Colaborador, '
           'Squad e Status da aba Ferias - pra adicionar uma pessoa nova ou um '
           'squad novo, edite a aba correspondente.')
trocas_9 = 0
for p in doc.paragraphs:
    if p.text.strip() == old_2_1:
        set_text(p, new_2_1)
        trocas_9 += 1
        break
print(f'(9) Aba Integrantes em 2.1: {trocas_9} paragrafo atualizado.')


# ============================================================
# 10) Insere subsecao "2.4 Cadastrar integrantes (datas de admissao
#     e ferias)" antes da secao 3
# ============================================================
ja_tem_24 = any('2.4 Cadastrar integrantes' in p.text for p in doc.paragraphs)
if not ja_tem_24:
    p_secao3 = None
    for p in doc.paragraphs:
        if p.text.strip() == '3. Gerar o relatorio':
            p_secao3 = p
            break
    if p_secao3 is not None:
        h3_tpl = find_template('Heading 3')
        first_tpl = find_template('First Paragraph')
        compact_tpl = find_template('Compact')

        blocks_24 = [
            (h3_tpl, '2.4 Cadastrar integrantes (datas de admissao e ferias)'),
            (first_tpl, 'A aba Integrantes alem de servir como cadastro de pessoas pros '
                        'dropdowns, tambem alimenta a secao "Controle de Vencimento de Ferias" '
                        'do relatorio (ver 6.2). Tem 4 colunas:'),
            (compact_tpl, 'Integrante - nome do colaborador (mesmo nome usado na aba Ferias).'),
            (compact_tpl, 'Squad - squad atual da pessoa.'),
            (compact_tpl, 'Data de inicio na AIR - data de admissao na empresa, no formato dd/mm/aaaa.'),
            (compact_tpl, 'Data das ultimas ferias - data fim do ultimo periodo de ferias tirado '
                          'pela pessoa. Pode ficar em branco se ela nunca tirou ferias (o app '
                          'calcula a partir da admissao).'),
            (first_tpl, 'Importante: a coluna "Data das ultimas ferias" e atualizada manualmente. '
                        'Toda vez que alguem volta de ferias, abra a aba Integrantes e atualize a '
                        'data dela. Sem essa atualizacao, o app continua avisando que a pessoa '
                        'precisa tirar ferias mesmo depois de ela ja ter tirado.'),
        ]
        inserted_24 = 0
        for tpl, txt in blocks_24:
            if tpl is None:
                continue
            new_p = clone_with_text(tpl, txt)
            p_secao3._element.addprevious(new_p)
            inserted_24 += 1
        print(f'(10) Subsecao 2.4 inserida ({inserted_24} paragrafos).')
    else:
        print('(10) Heading "3. Gerar o relatorio" nao encontrado, pulando.')
else:
    print('(10) Subsecao 2.4 ja existe, pulando.')


# ============================================================
# 11) Renumera 6.2-6.4 e insere "6.2 Controle de Vencimento de
#     Ferias" antes da nova 6.3
# ============================================================
ja_tem_62_venc = any('6.2 Controle de Vencimento' in p.text for p in doc.paragraphs)
if not ja_tem_62_venc:
    # 1. Renomeia 6.2 -> 6.3, 6.3 -> 6.4, 6.4 -> 6.5 (de tras pra frente
    #    pra nao colidir nomes)
    renames_6 = [
        ('6.4 Rodape de autoria',  '6.5 Rodape de autoria'),
        ('6.3 Gantt',              '6.4 Gantt'),
        ('6.2 Cronograma detalhado','6.3 Cronograma detalhado'),
    ]
    for old, new in renames_6:
        for p in doc.paragraphs:
            if p.text.strip() == old:
                set_text(p, new)
                break

    # 2. Atualiza o "tem 4 secoes" pra "tem 5 secoes" no paragrafo introdutorio
    for p in doc.paragraphs:
        if p.text.strip() == 'O HTML e o PDF tem 4 secoes:':
            set_text(p, 'O HTML e o PDF tem 5 secoes:')
            break

    # 3. Insere nova "6.2 Controle de Vencimento de Ferias" antes da
    #    nova 6.3 Cronograma detalhado
    p_63 = None
    for p in doc.paragraphs:
        if p.text.strip() == '6.3 Cronograma detalhado':
            p_63 = p
            break
    if p_63 is not None:
        h3_tpl = find_template('Heading 3')
        first_tpl = find_template('First Paragraph')
        compact_tpl = find_template('Compact')

        blocks_62 = [
            (h3_tpl, '6.2 Controle de Vencimento de Ferias'),
            (first_tpl, 'Aparece logo apos o Dashboard mensal. Calcula 1o e 2o vencimento de '
                        'ferias por colaborador (CLT 12 e 24 meses) e destaca quem precisa tirar '
                        'ferias antes de vencer. Janela de alerta: 6 meses antes de cada vencimento.'),
            (first_tpl, 'A secao tem 3 sub-blocos coloridos:'),
            (compact_tpl, 'CRITICO (vermelho) - pessoas com o 2o vencimento proximo (<= 6 meses) '
                          'ou ja vencido. Quem ja passou do 2o vencimento esta com ferias EM '
                          'DOBRO: pela CLT a empresa paga em duplicidade pelo periodo nao '
                          'tirado. Acao: programar ferias com urgencia.'),
            (compact_tpl, 'ATENCAO (laranja) - pessoas com o 1o vencimento proximo ou ja '
                          'atingido. Ja tem direito a tirar ferias mas ainda nao tirou. '
                          'Acao: planejar.'),
            (compact_tpl, 'Dados incompletos (cinza) - pessoas sem "Data de inicio na AIR" '
                          'preenchida na aba Integrantes. O app nao consegue calcular '
                          'vencimento sem essa data. Acao: abrir a planilha e atualizar.'),
            (first_tpl, 'Quem esta em dia (mais de 6 meses pra qualquer vencimento) NAO '
                        'aparece na secao -- ela fica focada so em quem precisa de atencao.'),
        ]
        inserted_62 = 0
        for tpl, txt in blocks_62:
            if tpl is None:
                continue
            new_p = clone_with_text(tpl, txt)
            p_63._element.addprevious(new_p)
            inserted_62 += 1
        print(f'(11) Secao 6.2 Vencimento + renumeracao 6.3-6.5: {inserted_62} paragrafos inseridos.')
    else:
        print('(11) Heading "6.3 Cronograma detalhado" nao encontrado pos-renumeracao, pulando.')
else:
    print('(11) Secao 6.2 Vencimento ja existe, pulando.')


# ============================================================
# Salva
# ============================================================
doc.save(str(DOCX))
print(f'OK: {DOCX} atualizado.')
